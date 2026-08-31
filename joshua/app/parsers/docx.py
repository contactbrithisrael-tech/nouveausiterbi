"""Parseur DOCX : les styles de titre deviennent des sections.

Word ne fournit pas de numéros de page (la pagination dépend du rendu). La
localisation exploitable est donc le titre de section, extrait des styles
« Heading N ». Une citation « Reglement.docx — section 3.2 » reste
vérifiable, ce qui est l'objectif.

Les tableaux sont extraits après le corps : ils contiennent souvent
l'information la plus dense (barèmes, listes de références) et les ignorer
amputerait le document d'une partie de sa substance.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

from app.ingestion.chunker import Bloc


def parser(chemin: Path) -> tuple[Iterator[Bloc], dict]:
    import docx  # python-docx, import tardif

    document = docx.Document(str(chemin))
    proprietes = document.core_properties
    infos = {
        "title": (proprietes.title or chemin.stem).strip(),
        "author": (proprietes.author or "").strip() or None,
        "year": proprietes.created.year if proprietes.created else None,
    }

    def generer() -> Iterator[Bloc]:
        section_courante: str | None = None
        tampon: list[str] = []

        def vider() -> Iterator[Bloc]:
            if tampon:
                yield Bloc(texte="\n\n".join(tampon), section=section_courante, title=infos["title"])
                tampon.clear()

        for paragraphe in document.paragraphs:
            texte = paragraphe.text.strip()
            if not texte:
                continue
            style = (paragraphe.style.name or "").lower()
            if style.startswith("heading") or style.startswith("titre"):
                # Un nouveau titre ferme la section précédente : c'est la
                # seule frontière structurelle fiable dans un DOCX.
                yield from vider()
                section_courante = texte[:256]
                continue
            tampon.append(texte)
        yield from vider()

        for i, tableau in enumerate(document.tables, start=1):
            lignes = []
            for ligne in tableau.rows:
                cellules = [c.text.strip() for c in ligne.cells if c.text.strip()]
                if cellules:
                    lignes.append(" | ".join(cellules))
            if lignes:
                yield Bloc(texte="\n".join(lignes), section=f"tableau {i}", title=infos["title"])

    return generer(), infos
